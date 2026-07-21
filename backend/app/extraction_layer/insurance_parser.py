import logging
import re
from collections import OrderedDict
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from app.extraction_layer.textract_client import TextractClient, TextractExtractionError
from app.schemas.analysis import ParsedDocument, UploadDescriptor

logger = logging.getLogger(__name__)


class DocumentExtractionError(ValueError):
    """Raised when an uploaded document cannot produce trustworthy text."""

SECTION_KEYWORDS = OrderedDict(
    {
        "insurance": r"\binsur\w*\b",
        "certificate": r"\bcertificate\b",
        "certificate holder": r"\bcertificate\s+holder\b",
        "coverage": r"\bcover\w*\b",
        "liability": r"\bliabilit\w*\b",
        "automobile": r"\bautomobile\b|\bauto\b",
        "workers compensation": r"\bworkers?\s+comp(?:ensation)?\b|\bWC\b",
        "employers liability": r"\bemployers?\s+liabilit\w*\b|\bE\.L\.\b",
        "indemnify": r"\bindemnif\w*\b",
        "indemnity": r"\bindemnit\w*\b",
        "additional insured": r"\badditional\s+insured\b|\bAI\b",
        "waiver of subrogation": r"\bwaiver\s+of\s+subrogation\b|\bsubrogation\s+waiv\w*\b",
        "umbrella": r"\bumbrella\b",
        "excess": r"\bexcess\b"
    }
)


class InsuranceDocumentParser:
    def __init__(self) -> None:
        self.textract = TextractClient()

    def parse(self, documents: list[UploadDescriptor]) -> list[ParsedDocument]:
        parsed_documents: list[ParsedDocument] = []

        for document in documents:
            file_name = document.file_name or document.document_id
            suffix = Path(file_name).suffix.lower()
            content = document.content or ""
            binary_payload = document.binary_payload

            if suffix == ".pdf" and binary_payload is not None:
                parsed = self._parse_pdf_bytes(binary_payload, document.document_id, file_name, document.document_type)
            elif suffix == ".docx" and binary_payload is not None:
                parsed = self._parse_docx_bytes(binary_payload, document.document_id, file_name, document.document_type)
            elif suffix in {".md", ".markdown"}:
                parsed = self._parse_markdown(content, document.document_id, file_name, document.document_type)
            else:
                parsed = self._parse_plain_text(content, document.document_id, file_name, document.document_type)

            parsed.description_box_lines = self._extract_description_box_lines(parsed.markdown)
            parsed.certificate_holder_text = self._extract_certificate_holder_text(parsed.markdown)
            sections, keywords = self._detect_insurance_sections(parsed.markdown, parsed.description_box_lines, parsed.certificate_holder_text)
            parsed.extracted_sections = sections
            parsed.matched_keywords = keywords
            parsed_documents.append(parsed)

            logger.debug("Parsed %s with matched keywords: %s", parsed.document_id, keywords)

        return parsed_documents

    def _parse_pdf_bytes(self, payload: bytes, document_id: str, file_name: str, document_type: str) -> ParsedDocument:
        embedded_text, page_count = self._extract_embedded_pdf_text(payload)
        if self._is_usable_text(embedded_text):
            return ParsedDocument(
                document_id=document_id,
                document_type=document_type,
                file_name=file_name,
                markdown=embedded_text,
                structured_json={
                    "name": file_name,
                    "format": "pdf",
                    "page_count": page_count,
                    "extraction_method": "embedded_pdf_text",
                },
                extracted_sections=[],
                extraction_method="embedded_pdf_text",
                extraction_confidence=self._score_text_quality(embedded_text, base=0.92),
            )

        try:
            result = self.textract.analyze_pdf(payload, file_name)
        except TextractExtractionError as exc:
            logger.exception("PDF extraction failed for %s", file_name)
            raise DocumentExtractionError(
                f"{file_name} could not be read reliably. Try a text-based PDF or a clearer scan."
            ) from exc

        return ParsedDocument(
            document_id=document_id,
            document_type=document_type,
            file_name=file_name,
            markdown=result.text,
            structured_json={
                "name": file_name,
                "format": "pdf",
                "page_count": result.page_count,
                "block_count": result.block_count,
                "textract_model_version": result.model_version,
                "extraction_method": "amazon_textract",
            },
            extracted_sections=[],
            extraction_method="amazon_textract",
            extraction_confidence=result.confidence,
        )

    def _extract_embedded_pdf_text(self, payload: bytes) -> tuple[str, int]:
        try:
            reader = PdfReader(BytesIO(payload), strict=False)
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception as exc:
                    raise DocumentExtractionError(
                        "Password-protected PDFs are not supported."
                    ) from exc
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(page.strip() for page in pages if page.strip()), len(reader.pages)
        except DocumentExtractionError:
            raise
        except Exception as exc:
            logger.info("Embedded PDF text extraction failed; trying Textract: %s", exc)
            return "", 0

    def _parse_docx_bytes(self, payload: bytes, document_id: str, file_name: str, document_type: str) -> ParsedDocument:
        try:
            document = Document(BytesIO(payload))
            blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        blocks.append(" | ".join(cells))
            text = "\n".join(blocks)
        except Exception as exc:
            logger.exception("DOCX extraction failed for %s", file_name)
            raise DocumentExtractionError(
                f"{file_name} could not be read as a Word document."
            ) from exc

        if not self._is_usable_text(text):
            raise DocumentExtractionError(f"{file_name} did not contain enough readable text.")

        return ParsedDocument(
            document_id=document_id,
            document_type=document_type,
            file_name=file_name,
            markdown=text,
            structured_json={"name": file_name, "format": "docx", "extraction_method": "python_docx"},
            extracted_sections=[],
            extraction_method="python_docx",
            extraction_confidence=self._score_text_quality(text, base=0.93),
        )

    def _is_usable_text(self, text: str) -> bool:
        compact = re.sub(r"\s+", " ", text).strip()
        if len(compact) < 40:
            return False
        printable_ratio = sum(character.isprintable() for character in compact) / len(compact)
        letter_ratio = sum(character.isalpha() for character in compact) / len(compact)
        return printable_ratio >= 0.95 and letter_ratio >= 0.25

    def _score_text_quality(self, text: str, base: float) -> float:
        compact = re.sub(r"\s+", " ", text).strip()
        if not compact:
            return 0.0
        printable_ratio = sum(character.isprintable() for character in compact) / len(compact)
        replacement_penalty = min(0.15, compact.count("\ufffd") * 0.02)
        length_adjustment = 0.02 if len(compact) >= 300 else 0.0
        return round(
            max(0.0, min(0.97, base + length_adjustment - (1 - printable_ratio) - replacement_penalty)),
            2,
        )

    def _parse_markdown(self, content: str, document_id: str, file_name: str, document_type: str) -> ParsedDocument:
        return ParsedDocument(
            document_id=document_id,
            document_type=document_type,
            file_name=file_name,
            markdown=content,
            structured_json={"name": file_name, "format": "markdown"},
            extracted_sections=[],
            extraction_method="markdown_text",
            extraction_confidence=self._score_text_quality(content, base=0.93),
        )

    def _parse_plain_text(self, content: str, document_id: str, file_name: str, document_type: str) -> ParsedDocument:
        return ParsedDocument(
            document_id=document_id,
            document_type=document_type,
            file_name=file_name,
            markdown=content.strip(),
            structured_json={"name": file_name, "format": "text", "content": content},
            extracted_sections=[],
            extraction_method="plain_text",
            extraction_confidence=self._score_text_quality(content, base=0.93),
        )

    def _detect_insurance_sections(
        self,
        markdown: str,
        description_box_lines: list[str],
        certificate_holder_text: str | None
    ) -> tuple[list[str], list[str]]:
        sections = self._split_into_sections(markdown)
        matched_sections: list[str] = []
        keyword_hits: list[str] = []

        for section in sections:
            matches = [name for name, pattern in SECTION_KEYWORDS.items() if re.search(pattern, section, re.IGNORECASE)]
            if matches:
                matched_sections.append(section)
                keyword_hits.extend(matches)

        if certificate_holder_text:
            synthetic_section = f"CERTIFICATE HOLDER\n{certificate_holder_text}"
            matched_sections.append(synthetic_section)
            keyword_hits.extend(["certificate", "certificate holder"])

        for line in description_box_lines:
            synthetic_line = f"DESCRIPTION OF OPERATIONS\n{line}"
            matches = [name for name, pattern in SECTION_KEYWORDS.items() if re.search(pattern, synthetic_line, re.IGNORECASE)]
            if matches:
                matched_sections.append(synthetic_line)
                keyword_hits.extend(matches)

        return matched_sections, list(OrderedDict.fromkeys(keyword_hits))

    def _split_into_sections(self, markdown: str) -> list[str]:
        lines = [line.strip() for line in markdown.splitlines() if line.strip()]
        if not lines:
            return []

        # PDF text extractors preserve lines more reliably than headings. Small,
        # overlapping windows keep each coverage label beside its own limits.
        if not any(line.startswith("#") for line in lines):
            return ["\n".join(lines[index:index + 4]) for index in range(len(lines))]

        sections: list[str] = []
        current: list[str] = []

        for line in lines:
            stripped = line.strip()
            is_heading = stripped.startswith("#") or stripped.isupper()
            if is_heading and current:
                sections.append("\n".join(current).strip())
                current = [line]
            else:
                current.append(line)

        if current:
            sections.append("\n".join(current).strip())

        if len(sections) <= 1:
            sections = [block.strip() for block in re.split(r"\n\s*\n", markdown) if block.strip()]

        return sections

    def _extract_certificate_holder_text(self, markdown: str) -> str | None:
        line_match = re.search(
            r"(?im)^certificate holder(?: name)?\s*:?\s*$\n([^\n]+)",
            markdown,
        )
        if line_match:
            return line_match.group(1).strip(" .:-")

        patterns = [
            r"CERTIFICATE HOLDER\s+(.*?)\s+AUTHORIZED REPRESENTATIVE",
            r"CERTIFICATE HOLDER\s+(.*?)\s+SHOULD ANY OF THE ABOVE DESCRIBED POLICIES",
        ]
        compact = re.sub(r"\s+", " ", markdown).strip()
        for pattern in patterns:
            match = re.search(pattern, compact, re.IGNORECASE)
            if match:
                return match.group(1).strip(" .:-")
        return None

    def _extract_description_box_lines(self, markdown: str) -> list[str]:
        start_markers = [
            "DESCRIPTION OF OPERATIONS / LOCATIONS / VEHICLES",
            "DESCRIPTION OF OPERATIONS",
        ]
        end_markers = [
            "SHOULD ANY OF THE ABOVE DESCRIBED POLICIES",
            "CERTIFICATE HOLDER",
            "AUTHORIZED REPRESENTATIVE",
        ]

        upper_markdown = markdown.upper()
        start_index = -1
        start_marker_text = ""
        for marker in start_markers:
            marker_index = upper_markdown.find(marker.upper())
            if marker_index != -1:
                start_index = marker_index
                start_marker_text = marker
                break

        if start_index == -1:
            return []

        content = markdown[start_index + len(start_marker_text):]
        end_index = len(content)
        upper_content = content.upper()
        for marker in end_markers:
            marker_index = upper_content.find(marker.upper())
            if marker_index != -1 and marker_index < end_index:
                end_index = marker_index

        block = content[:end_index]
        if not block.strip():
            return []

        candidates: list[str] = []
        for line in re.split(r"[\r\n]+", block):
            cleaned = re.sub(r"\s+", " ", line).strip(" .:-")
            if len(cleaned) < 4:
                continue
            if cleaned.upper() == cleaned and len(cleaned.split()) > 6:
                continue
            if cleaned.lower().startswith("attach acord 101"):
                continue
            candidates.append(cleaned)

        if not candidates:
            compact = re.sub(r"\s+", " ", block).strip()
            known_line_patterns = [
                r"(Liquor Liability\s+\$[\d,]+(?:\.\d+)?\s+\w+\s+\w+)",
                r"(Pollution Liability\s+\$[\d,]+(?:\.\d+)?\s+\w+\s+\w+)",
                r"(Aviation/UAV Liability\s+\$[\d,]+(?:\.\d+)?\s+\w+\s+\w+)",
                r"(Marine Liability\s+\$[\d,]+(?:\.\d+)?\s+\w+\s+\w+)",
                r"(Professional Liability\s+\$[\d,]+(?:\.\d+)?)",
                r"(Demolition, Roofing or Electrical Contractors\s+\$[\d,]+(?:\.\d+)?/[A-Za-z]+)",
                r"(Transportation Svc > \d+ Passengers\s+\$[\d,]+(?:\.\d+)?/[A-Za-z]+)",
            ]
            for pattern in known_line_patterns:
                for match in re.findall(pattern, compact, re.IGNORECASE):
                    candidates.append(match.strip())

        unique_candidates: list[str] = []
        for candidate in candidates:
            if candidate not in unique_candidates:
                unique_candidates.append(candidate)
        return unique_candidates
